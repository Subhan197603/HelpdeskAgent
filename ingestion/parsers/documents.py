"""Bounded parser adapters for approved acquisition content types."""

import csv
import io
import json
import re
from dataclasses import dataclass
from html import unescape
from typing import Protocol

import fitz  # type: ignore[import-untyped]
from bs4 import BeautifulSoup, Tag
from docx import Document


class ParsingError(RuntimeError):
    """An acquired document cannot be safely parsed."""


@dataclass(frozen=True, slots=True)
class StoredFile:
    content: bytes
    filename: str
    content_type: str


@dataclass(frozen=True, slots=True)
class ParsedSection:
    heading_path: tuple[str, ...]
    text: str
    page_number: int | None
    anchor: str | None
    content_type: str
    table_markdown: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    sections: tuple[ParsedSection, ...]
    page_count: int | None
    metadata: dict[str, str]

    @property
    def normalized_text(self) -> str:
        return "\n\n".join(section.text for section in self.sections if section.text.strip())


class DocumentParser(Protocol):
    name: str
    version: str
    supported_types: frozenset[str]

    async def parse(self, file: StoredFile) -> ParsedDocument: ...


def _clean(value: str) -> str:
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[ \t]+", " ", normalized).strip()


class PlainTextParser:
    name = "plain-structured-text"
    version = "1"
    supported_types = frozenset({"text/plain", "text/markdown", "text/csv", "application/json"})

    async def parse(self, file: StoredFile) -> ParsedDocument:
        try:
            value = file.content.decode("utf-8")
        except UnicodeDecodeError as error:
            raise ParsingError("TEXT_NOT_UTF8") from error
        if file.content_type == "application/json":
            try:
                value = json.dumps(json.loads(value), ensure_ascii=False, indent=2, sort_keys=True)
            except json.JSONDecodeError as error:
                raise ParsingError("INVALID_JSON") from error
        if file.content_type == "text/csv":
            value = _csv_markdown(value)
        sections = _markdown_sections(value, file.content_type)
        if not sections:
            raise ParsingError("EMPTY_EXTRACTION")
        return ParsedDocument(tuple(sections), None, {"encoding": "utf-8"})


class HtmlDocumentParser:
    name = "beautifulsoup-html"
    version = "1"
    supported_types = frozenset({"text/html"})

    async def parse(self, file: StoredFile) -> ParsedDocument:
        soup = BeautifulSoup(file.content, "html.parser")
        for element in soup(["script", "style", "noscript", "template"]):
            element.decompose()
        headings: list[str] = []
        sections: list[ParsedSection] = []
        body = soup.body or soup
        for node in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "li", "table"]):
            if not isinstance(node, Tag):
                continue
            if node.name and node.name.startswith("h"):
                level = int(node.name[1])
                headings = headings[: level - 1] + [_clean(node.get_text(" ", strip=True))]
                continue
            table = _html_table(node) if node.name == "table" else None
            text = table or _clean(unescape(node.get_text(" ", strip=True)))
            if text:
                sections.append(
                    ParsedSection(
                        tuple(headings),
                        text,
                        None,
                        str(node.get("id")) if node.get("id") else None,
                        "table" if table else "text",
                        table,
                    )
                )
        if not sections:
            text = _clean(body.get_text("\n", strip=True))
            if text:
                sections.append(ParsedSection((), text, None, None, "text"))
        if not sections:
            raise ParsingError("EMPTY_EXTRACTION")
        title = _clean(soup.title.get_text(" ")) if soup.title else ""
        return ParsedDocument(tuple(sections), None, {"title": title} if title else {})


class PdfDocumentParser:
    name = "pymupdf"
    version = "1"
    supported_types = frozenset({"application/pdf"})

    async def parse(self, file: StoredFile) -> ParsedDocument:
        try:
            document = fitz.open(stream=file.content, filetype="pdf")
        except Exception as error:
            raise ParsingError("INVALID_PDF") from error
        sections: list[ParsedSection] = []
        try:
            for number, page in enumerate(document, start=1):
                text = _clean(page.get_text("text"))
                if text:
                    sections.append(ParsedSection((), text, number, None, "text"))
            metadata = {
                str(key): str(value) for key, value in (document.metadata or {}).items() if value
            }
            page_count = document.page_count
        finally:
            document.close()
        if not sections:
            raise ParsingError("EMPTY_EXTRACTION")
        return ParsedDocument(tuple(sections), page_count, metadata)


class DocxDocumentParser:
    name = "python-docx"
    version = "1"
    supported_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    async def parse(self, file: StoredFile) -> ParsedDocument:
        try:
            document = Document(io.BytesIO(file.content))
        except Exception as error:
            raise ParsingError("INVALID_DOCX") from error
        headings: list[str] = []
        sections: list[ParsedSection] = []
        for paragraph in document.paragraphs:
            text = _clean(paragraph.text)
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else ""
            match = re.match(r"Heading (\d+)", style)
            if match:
                level = int(match.group(1))
                headings = headings[: level - 1] + [text]
            else:
                sections.append(ParsedSection(tuple(headings), text, None, None, "text"))
        for table in document.tables:
            rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
            markdown = _rows_markdown(rows)
            if markdown:
                sections.append(
                    ParsedSection(tuple(headings), markdown, None, None, "table", markdown)
                )
        if not sections:
            raise ParsingError("EMPTY_EXTRACTION")
        return ParsedDocument(tuple(sections), None, {})


class ParserRegistry:
    def __init__(self, parsers: tuple[DocumentParser, ...] | None = None) -> None:
        configured = parsers or (
            PlainTextParser(),
            HtmlDocumentParser(),
            PdfDocumentParser(),
            DocxDocumentParser(),
        )
        self._parsers = {
            content_type: parser for parser in configured for content_type in parser.supported_types
        }

    def require(self, content_type: str) -> DocumentParser:
        parser = self._parsers.get(content_type.lower())
        if parser is None:
            raise ParsingError("UNSUPPORTED_CONTENT_TYPE")
        return parser


def _markdown_sections(value: str, content_type: str) -> list[ParsedSection]:
    headings: list[str] = []
    buffer: list[str] = []
    sections: list[ParsedSection] = []

    def flush() -> None:
        text = _clean("\n".join(buffer))
        if text:
            sections.append(ParsedSection(tuple(headings), text, None, None, content_type))
        buffer.clear()

    for line in value.splitlines():
        match = re.match(r"^(#{1,6})\s+(.+)$", line) if content_type == "text/markdown" else None
        if match:
            flush()
            level = len(match.group(1))
            headings = headings[: level - 1] + [_clean(match.group(2))]
        elif not line.strip() and buffer:
            flush()
        else:
            buffer.append(line)
    flush()
    return sections


def _csv_markdown(value: str) -> str:
    try:
        rows = list(csv.reader(io.StringIO(value)))
    except csv.Error as error:
        raise ParsingError("INVALID_CSV") from error
    return _rows_markdown(rows)


def _rows_markdown(rows: list[list[str]]) -> str:
    if not rows or not rows[0]:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = [
        "| " + " | ".join(normalized[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def _html_table(table: Tag) -> str | None:
    rows = [
        [_clean(cell.get_text(" ", strip=True)) for cell in row.find_all(["th", "td"])]
        for row in table.find_all("tr")
    ]
    return _rows_markdown(rows) or None
