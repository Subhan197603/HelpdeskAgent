"""Parser, chunker, embedding, validation, and publication-policy tests."""

import asyncio
import io
from uuid import UUID

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document
from pydantic import ValidationError as PydanticValidationError

from apps.api.app.core.context import RequestContext
from apps.api.app.identity.authorization import AuthorizationService, Permission
from apps.worker.worker.settings import WorkerSettings
from ingestion.chunkers import ChunkingConfig, SemanticChunker
from ingestion.embeddings import DeterministicEmbeddingProvider, HttpEmbeddingProvider
from ingestion.parsers import ParsedDocument, ParsedSection, ParserRegistry, StoredFile
from ingestion.validation import ValidationContext, validate_corpus

TENANT = UUID("20000000-0000-0000-0000-000000000001")
USER = UUID("22000000-0000-0000-0000-000000000007")


def _parse(content: bytes, filename: str, content_type: str) -> ParsedDocument:
    parser = ParserRegistry().require(content_type)
    return asyncio.run(parser.parse(StoredFile(content, filename, content_type)))


def test_markdown_parser_preserves_heading_lineage_and_normalized_text() -> None:
    parsed = _parse(
        b"# Install\n\nRead the warning.\n\n## Steps\n\n1. Configure.\n2. Verify.",
        "guide.md",
        "text/markdown",
    )
    assert parsed.sections[0].heading_path == ("Install",)
    assert parsed.sections[-1].heading_path == ("Install", "Steps")
    assert "Configure" in parsed.normalized_text


def test_html_parser_removes_executable_content_and_preserves_table() -> None:
    parsed = _parse(
        b"<html><head><title>Policy</title><script>secret()</script></head>"
        b"<body><h1>Scope</h1><p>Employees</p><table><tr><th>A</th></tr>"
        b"<tr><td>B</td></tr></table></body></html>",
        "policy.html",
        "text/html",
    )
    assert "secret" not in parsed.normalized_text
    assert parsed.sections[0].heading_path == ("Scope",)
    assert parsed.sections[-1].table_markdown == "| A |\n| --- |\n| B |"


def test_pdf_and_docx_parser_adapters_extract_text() -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Approved PDF procedure")
    pdf_bytes = pdf.tobytes()
    pdf.close()
    assert (
        "Approved PDF procedure"
        in _parse(pdf_bytes, "guide.pdf", "application/pdf").normalized_text
    )

    document = Document()
    document.add_heading("Procedure", level=1)
    document.add_paragraph("Complete the approved step.")
    buffer = io.BytesIO()
    document.save(buffer)
    docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    parsed = _parse(buffer.getvalue(), "guide.docx", docx_type)
    assert parsed.sections[0].heading_path == ("Procedure",)


def test_chunking_is_deterministic_and_embedding_input_contains_title_and_heading() -> None:
    parsed = ParsedDocument(
        (
            ParsedSection(("Install",), "Run the supported installer.", 1, "install", "text"),
            ParsedSection(("Install",), "Verify the result.", 1, None, "text"),
        ),
        1,
        {},
    )
    chunker = SemanticChunker(ChunkingConfig(100, 150, 50, 10))
    first = chunker.chunk("Deployment Guide", parsed)
    second = chunker.chunk("Deployment Guide", parsed)
    assert first == second
    assert first[0].embedding_input.startswith("Title: Deployment Guide\nHeading: Install")
    assert first[0].content_hash == second[0].content_hash


def test_chunking_configuration_and_embedding_inputs_are_version_sensitive() -> None:
    assert ChunkingConfig(100, 150, 50, 10).digest != ChunkingConfig(110, 150, 50, 10).digest
    provider = DeterministicEmbeddingProvider(dimension=8)
    first = asyncio.run(provider.embed_batch(["Title A"]))[0]
    assert first == asyncio.run(provider.embed_batch(["Title A"]))[0]
    assert first != asyncio.run(provider.embed_batch(["Title B"]))[0]
    assert len(first) == 8


def test_validation_fails_incomplete_embeddings_and_warns_before_human_approval() -> None:
    parsed = ParsedDocument((ParsedSection((), "Useful text", None, None, "text"),), None, {})
    chunks = SemanticChunker(ChunkingConfig(100, 150, 50, 10)).chunk("Policy", parsed)
    failed = validate_corpus(
        parsed,
        chunks,
        ValidationContext("en", True, True, True, "DRAFT", None, None, 0),
    )
    assert failed.status == "FAILED"
    assert "INCOMPLETE_EMBEDDINGS" in failed.failures
    warning = validate_corpus(
        parsed,
        chunks,
        ValidationContext("en", True, True, True, "DRAFT", None, None, len(chunks)),
    )
    assert warning.status == "WARNING"
    assert "DOCUMENT_NOT_APPROVED" in warning.warnings


def test_http_embedding_provider_requires_https() -> None:
    with pytest.raises(ValueError, match="HTTPS"):
        HttpEmbeddingProvider("http://provider.invalid/embeddings", "secret", "MODEL")


def test_production_rejects_deterministic_embedding_provider() -> None:
    with pytest.raises(PydanticValidationError, match="embedding provider"):
        WorkerSettings.model_validate(
            {
                "app_env": "production",
                "json_logs": True,
                "worker_database_url": (
                    "postgresql+psycopg://worker:strong-password@database.internal/helpdesk"
                ),
                "smtp_host": "smtp.example.invalid",
                "smtp_starttls": True,
                "object_storage_access_key": "production-access",
                "object_storage_secret_key": "production-secret",
                "object_storage_server_side_encryption": "AES256",
                "embedding_provider_mode": "deterministic",
            }
        )


@pytest.mark.parametrize(
    ("role", "allowed", "denied"),
    [
        (
            "KNOWLEDGE_AUTHOR",
            Permission.KNOWLEDGE_DOCUMENT_READ_ADMIN,
            Permission.KNOWLEDGE_DOCUMENT_PUBLISH,
        ),
        (
            "KNOWLEDGE_APPROVER",
            Permission.KNOWLEDGE_DOCUMENT_PUBLISH,
            Permission.KNOWLEDGE_MANIFEST_IMPORT,
        ),
        (
            "CUSTOMER",
            Permission.CATALOG_SERVICE_READ,
            Permission.KNOWLEDGE_DOCUMENT_READ_ADMIN,
        ),
    ],
)
def test_document_publication_roles_are_separated(
    role: str, allowed: Permission, denied: Permission
) -> None:
    context = RequestContext(
        tenant_id=TENANT,
        user_id=USER,
        external_subject="processing-role-test",
        roles=frozenset({role}),
        support_group_ids=frozenset(),
        business_unit_id=None,
        correlation_id=str(TENANT),
        request_id="processing-role-test",
    )
    authorization = AuthorizationService()
    assert authorization.is_allowed(context, allowed)
    assert not authorization.is_allowed(context, denied)
